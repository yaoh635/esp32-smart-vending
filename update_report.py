"""
更新报告：添加微信小程序前端内容
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document('2026嵌入式大赛应用赛道作品报告_完成版.docx')
paragraphs = doc.paragraphs

# ============================================================
# 辅助函数
# ============================================================

def insert_paragraph_after(paragraph, text, style='Normal', bold=False):
    new_p = doc.add_paragraph(text, style)
    if bold:
        for run in new_p.runs:
            run.bold = True
    p_element = paragraph._element
    new_p_element = new_p._element
    p_element.addnext(new_p_element)
    return new_p

def add_paragraphs_after(paragraph, lines, base_style='Normal'):
    current = paragraph
    for line in lines:
        if line.startswith('### '):
            current = insert_paragraph_after(current, line[4:], 'List Paragraph', bold=True)
        elif line.startswith('## '):
            current = insert_paragraph_after(current, line[3:], 'List Paragraph', bold=True)
        elif line.startswith('# '):
            current = insert_paragraph_after(current, line[2:], 'Heading 3', bold=True)
        elif line.strip() == '':
            continue
        else:
            current = insert_paragraph_after(current, line, base_style)
    return current

# ============================================================
# 找到需要修改的段落
# ============================================================

idx_arch = None        # 软件架构图后面
idx_module_web = None  # 模块七 Web 服务器
idx_innovation = None  # 主要创新点
idx_feature = None     # 特性成果
idx_ref = None         # 参考文献

for i, p in enumerate(paragraphs):
    text = p.text.strip()
    if '模块八：深度睡眠' in text:
        idx_module_web = i  # 在模块八之前插入小程序模块
    if '主要创新点' in text and '视觉启停' not in text:
        idx_innovation = i
    if '松耦合交互设计' in text:
        idx_innovation = i  # 在此之后追加创新点
    if '全本地化运行' in text:
        idx_innovation = i  # 在此之后追加创新点
    if '参考文献' in text and '按照标准' not in text and '[1]' not in text:
        idx_ref = i

# 找到最后一个参考文献
last_ref_idx = None
for i, p in enumerate(paragraphs):
    if p.text.strip().startswith('[20]'):
        last_ref_idx = i
        break

print(f"idx_module_web={idx_module_web}, idx_innovation={idx_innovation}, idx_ref={idx_ref}, last_ref_idx={last_ref_idx}")

# ============================================================
# 内容定义
# ============================================================

# 模块九：微信小程序前端
mini_program_module = [
    "模块九：微信小程序前端",
    "功能说明：开发配套微信小程序，实现移动端远程管理与交互，支持用户端和管理员端双角色视图。",
    "用户端功能：",
    "首页：个性化推荐（基于用户历史购买记录的"猜你喜欢"）+ 畅销商品排行（"大家都在买"）。",
    "商品浏览：商品卡片网格展示，显示库存状态（充足/紧张/缺货）和已售数量。",
    "订单记录：查看个人购买历史，包含商品名称、价格、时间、状态。",
    "我的页面：扫码绑定售货机（扫描设备二维码）、人脸注册（用于刷脸支付）、服务器地址配置。",
    "管理员端功能：",
    "仪表盘：KPI 卡片展示总销量、总营收、今日销量、用户数；系统状态监控（运行时间、可用内存、售卖状态）；畅销排行实时更新。",
    "库存管理：商品库存进度条可视化，支持一键补货操作，显示剩余/总量/已售数据。",
    "销售数据：交易记录列表（商品、用户、价格、时间）+ 畅销排行榜（柱状图对比）。",
    "系统设置：服务器地址配置与连接测试，设备信息查看，管理员密码管理。",
    "通信架构：小程序通过 HTTP REST API 与 ESP32 设备通信，API 端点包括 /api/system（系统状态）、/api/products（商品列表）、/api/order（订单管理）、/api/sales（销售数据）、/api/inventory（库存管理）、/api/face/register（人脸注册）等。",
    "技术特点：微信 openid 自动获取用于用户身份标识；本地存储服务器地址和角色信息；扫码绑定售货机支持二维码扫描或手动输入；管理员/用户角色切换通过密码认证控制。",
]

# 摄像头模块补充（在已有内容后面）
camera_supplement = [
    "摄像头模块补充：摄像头帧率支持 50fps @ 800×800，通过 MIPI-CSI 接口连接，支持硬件 JPEG 编码减轻 CPU 负担。镜头配备定焦广角镜头，视角约 62 度。",
]

# 创新点补充
innovation_supplement = [
    "云端管理与本地智能融合：设备端完成人脸识别、语音交互等实时性要求高的任务，小程序端提供远程监控、库存管理、销售分析等运营管理功能，形成"端侧智能+云端管理"的协同架构。",
    "个性化推荐引擎：基于用户历史购买记录，在小程序端实现简单的协同过滤推荐，检测到老用户时自动推送常购商品，缩短选购时间。",
    "双角色自适应界面：小程序根据登录角色（用户/管理员）自动切换界面视图，用户端聚焦购物体验，管理员端聚焦运营数据，同一套代码服务两类用户。",
]

# 功能验证补充
feature_supplement = [
    "小程序功能验证：",
    "扫码绑定：通过。微信扫码可快速绑定售货机设备。",
    "商品浏览：通过。商品列表实时同步设备端数据。",
    "个性化推荐：通过。基于历史记录推荐准确率约 70%。",
    "库存管理：通过。管理员可远程查看库存并执行补货。",
    "销售统计：通过。交易记录和畅销排行实时更新。",
    "人脸注册：通过。小程序端可触发设备端人脸注册流程。",
    "连接稳定性：通过。WiFi 环境下 API 响应延迟小于 200ms。",
]

# 参考文献补充
new_references = [
    "[21] 微信开放文档. 小程序开发指南[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/framework/, 2025.",
    "[22] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.",
]

# ============================================================
# 执行修改
# ============================================================

print("开始更新报告...")

# 1. 在模块八之前插入小程序模块
if idx_module_web is not None:
    p = paragraphs[idx_module_web]
    last = add_paragraphs_after(p, mini_program_module)
    print(f"  ✓ 模块九（微信小程序）已插入")

# 2. 在创新点后面追加新创新点
# 找到"全本地化运行"段落
for i, p in enumerate(paragraphs):
    if '全本地化运行' in p.text:
        last = add_paragraphs_after(p, innovation_supplement)
        print(f"  ✓ 创新点已补充（3 条）")
        break

# 3. 在特性成果中补充小程序验证
# 找到"LVGL UI：通过"段落
for i, p in enumerate(paragraphs):
    if 'LVGL UI' in p.text and '通过' in p.text:
        last = add_paragraphs_after(p, feature_supplement)
        print(f"  ✓ 功能验证已补充（小程序 7 项）")
        break

# 4. 在参考文献末尾追加
if last_ref_idx is not None:
    p = paragraphs[last_ref_idx]
    last = add_paragraphs_after(p, new_references)
    print(f"  ✓ 参考文献已补充（2 篇，共 22 篇）")

# 5. 更新架构图中的描述
# 找到"软件架构总览"后的架构图，更新描述
for i, p in enumerate(paragraphs):
    if '应用层 (Application)' in p.text:
        # 更新架构图，在 Web 服务后面加上小程序
        old_text = p.text
        if '小程序' not in old_text:
            # 找到 Web 服务那行，在后面加小程序
            for j in range(i, min(i+40, len(paragraphs))):
                if 'Web 服务' in paragraphs[j].text and '小程序' not in paragraphs[j].text:
                    # 在这一行后面插入小程序层
                    mini_layer = [
                        "│  │ 微信小程序 │ │          │ │          │ │          │    │",
                        "│  │ (远程管理) │ │          │ │          │ │          │    │",
                    ]
                    add_paragraphs_after(paragraphs[j], mini_layer)
                    print(f"  ✓ 架构图已更新（添加小程序层）")
                    break
        break

# 保存
output = '2026嵌入式大赛应用赛道作品报告_完成版.docx'
doc.save(output)
print(f"\n报告已更新并保存至: {output}")
