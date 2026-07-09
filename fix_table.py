"""
修复性能指标表格样式：白色背景 + 统一字体
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document('2026嵌入式大赛应用赛道作品报告_完成版.docx')

table = doc.tables[0]

print(f'表格: {len(table.rows)} 行 x {len(table.columns)} 列')

for row_idx, row in enumerate(table.rows):
    for cell in row.cells:
        # 设置单元格背景为白色
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            tc.insert(0, tcPr)

        # 删除已有的背景色
        for shd in tcPr.findall(qn('w:shd')):
            tcPr.remove(shd)

        # 设置白色背景
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF" w:val="clear"/>')
        tcPr.append(shd)

        # 设置单元格边框为浅灰色
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'</w:tcBorders>'
        )
        for old_borders in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old_borders)
        tcPr.append(borders)

        # 设置字体样式
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10.5)  # 五号字体
                run.font.name = '宋体'
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色字体
                # 设置中文字体
                rPr = run._element.find(qn('w:rPr'))
                if rPr is None:
                    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                    run._element.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), '宋体')

# 保存
output = '2026嵌入式大赛应用赛道作品报告_完成版.docx'
doc.save(output)
print(f'表格样式已修复，保存至: {output}')
